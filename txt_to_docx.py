import re
import sys
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==============================================================================
# 1. FUNÇÕES DE LIMPEZA E FORMATAÇÃO
# ==============================================================================

def clean_and_format_text(text):
    """
    Limpa tags, traduz LaTeX e normaliza texto.
    """
    # Remove tags [source]
    text = re.sub(r'\\', '', text).strip()

    # Substituições Visuais
    replacements = [
        (r'\\?sqrt', '√'), (r'\\?approx', '≈'), (r'\\?neq', '≠'),
        (r'\\?le\b', '≤'), (r'\\?ge\b', '≥'), (r'\\?times', '×'),
        (r'\\?cdot', '·'), (r'\\?infty', '∞'), (r'\\?rightarrow', '→'),
        (r'\\?pm', '±'), (r'\\?sum', '∑'),
        (r'\\?alpha', 'α'), (r'\\?beta', 'β'), (r'\\?theta', 'θ'),
        (r'\\?mu', 'μ'), (r'\\?sigma', 'σ'), (r'\\?lambda', 'λ'), (r'\\?pi', 'π'),
        (r'\^2', '²'), (r'\^3', '³'),
        (r'_E', 'ᴇ'), (r'_{WE}', 'ᴡᴇ'), (r'_M', 'ᴍ'),
        (r'\\?frac\{1\}\{2\}', '½'),
        (r'\\?frac', ''), 
        (r'\\?textbf', ''), (r'\\?textit', ''),
    ]
    
    clean = text
    for pattern, replacement in replacements:
        clean = re.sub(pattern, replacement, clean)
    
    # Limpeza final segura
    clean = clean.replace('\\', '').replace('{', '').replace('}', '')
    clean = re.sub(r'\s+', ' ', clean)
    
    return clean.strip()

def add_formatted_run(paragraph, text, force_bold=False):
    """
    Adiciona texto ao parágrafo processando negrito (**...**).
    """
    # Normaliza textbf para **
    text = re.sub(r'(?i)\\?textbf\{(.*?)\}', r'**\1**', text)
    
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            clean_txt = clean_and_format_text(part[2:-2])
            if clean_txt:
                run = paragraph.add_run(clean_txt)
                run.bold = True
                run.font.name = "Calibri"
        else:
            clean_txt = clean_and_format_text(part)
            if clean_txt:
                run = paragraph.add_run(clean_txt)
                run.bold = force_bold
                run.font.name = "Calibri"

# ==============================================================================
# 2. PROCESSADOR DE TABELAS
# ==============================================================================

def process_table_block(doc, raw_lines):
    rows = []
    for line in raw_lines:
        clean = re.sub(r'\\', '', line).strip()
        if any(x in clean for x in ['begin{table}', 'begin{tabular}', 'end{table}', 'end{tabular}', 'caption', 'hline']):
            continue
        if not clean: continue

        clean = clean.replace('\\', '')
        cols = [c.strip() for c in clean.split('&')]
        if any(cols): rows.append(cols)

    if not rows: return

    # Tabela com bordas
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                p = row.cells[j].paragraphs[0]
                add_formatted_run(p, cell_text, force_bold=(i == 0))
                # Ajuste fino de espaçamento na tabela
                p.paragraph_format.space_after = Pt(0)
    
    doc.add_paragraph("")

# ==============================================================================
# 3. LÓGICA PRINCIPAL (COM CONTROLE MANUAL DE LISTAS)
# ==============================================================================

def txt_to_docx(txt_path, docx_path):
    document = Document()
    
    with open(txt_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    # Estrutura para controlar listas aninhadas
    # Cada item na pilha é um dicionário: {'type': 'ul'/'ol', 'count': 1}
    list_stack = [] 
    
    in_table = False
    table_buffer = []

    for raw_line in lines:
        line_content = re.sub(r'\\', '', raw_line).strip()
        if not line_content: continue

        # --- TABELAS ---
        if 'begin{tabular}' in line_content or 'begin{table}' in line_content:
            in_table = True
            table_buffer = []
            continue
        if in_table:
            table_buffer.append(line_content)
            if 'end{tabular}' in line_content or 'end{table}' in line_content:
                in_table = False
                process_table_block(document, table_buffer)
            continue

        # --- CONTROLE DE INÍCIO/FIM DE LISTAS ---
        # Enumerate (Numerada)
        if 'begin{enumerate}' in line_content:
            list_stack.append({'type': 'ol', 'count': 1})
            continue
        # Itemize (Bolinhas)
        if 'begin{itemize}' in line_content:
            list_stack.append({'type': 'ul', 'count': 0})
            continue
        
        # Fim de lista
        if 'end{itemize}' in line_content or 'end{enumerate}' in line_content:
            if list_stack: list_stack.pop()
            continue

        # --- ITENS DA LISTA ---
        if re.search(r'^\\?item\s', line_content):
            # Limpa o comando \item
            text_part = re.sub(r'^\\?item\s', '', line_content).strip()
            
            # Recupera dados da lista atual (topo da pilha)
            if list_stack:
                current_list = list_stack[-1]
                level = len(list_stack)
                
                # Define o prefixo (Bolinha ou Número)
                if current_list['type'] == 'ol':
                    prefix = f"{current_list['count']}. "
                    current_list['count'] += 1 # Incrementa contador para o próximo
                else:
                    prefix = "•  " # Bolinha clássica
                
                # Cria parágrafo normal (não usa estilo de lista do Word para evitar bug de contagem)
                p = document.add_paragraph()
                
                # Adiciona o prefixo (1. ou •)
                run_prefix = p.add_run(prefix)
                run_prefix.font.name = "Calibri"
                run_prefix.bold = True # Número/Bolinha em negrito fica melhor
                
                # Adiciona o texto formatado
                add_formatted_run(p, text_part)

                # --- MÁGICA DO ESPAÇAMENTO (TABULAÇÃO) ---
                # Recuo à esquerda (joga tudo para direita)
                # 0.75 cm por nível de identação
                indent_size = Cm(0.75) * level 
                p.paragraph_format.left_indent = indent_size
                
                # Recuo da primeira linha (traz o número/bolinha de volta para a esquerda)
                # Cria o efeito de "Hanging Indent" (texto alinhado, número solto)
                p.paragraph_format.first_line_indent = Cm(-0.63)
                
                # Remove espaçamento excessivo entre itens
                p.paragraph_format.space_after = Pt(2)
            else:
                # Caso raro: item fora de lista
                p = document.add_paragraph()
                add_formatted_run(p, text_part)
            continue

        # --- TÍTULOS ---
        if line_content.startswith('#'):
            level = line_content.count('#')
            title = line_content.replace('#', '').strip()
            p = document.add_heading(title, level=min(level, 9))
            for r in p.runs: r.font.color.rgb = RGBColor(0,0,0)
            continue

        # --- FÓRMULAS $$ ---
        if '$$' in line_content:
            math = clean_and_format_text(line_content.replace('$$', ''))
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(math)
            run.italic = True
            run.font.name = "Cambria Math"
            run.font.size = Pt(12)
            continue

        # --- PARÁGRAFOS COMUNS ---
        # Texto dentro de lista mas sem \item (continuação)
        if list_stack:
            p = document.add_paragraph()
            add_formatted_run(p, line_content)
            # Alinha com o texto do item acima
            level = len(list_stack)
            p.paragraph_format.left_indent = Cm(0.75) * level
            p.paragraph_format.space_after = Pt(2)
        else:
            # Texto fora de lista
            p = document.add_paragraph()
            add_formatted_run(p, line_content)
            p.paragraph_format.space_after = Pt(6)

    document.save(docx_path)
    print(f"Sucesso! DOCX gerado: {docx_path}")

# ==============================================================================
if __name__ == "__main__":
    target = sys.argv[1] if len(sys.argv) > 1 else "mineracao.txt"
    if os.path.exists(target):
        try:
            txt_to_docx(target, target.replace('.txt', '.docx'))
        except Exception as e:
            print(f"Erro: {e}")
    else:
        print("Arquivo não encontrado.")